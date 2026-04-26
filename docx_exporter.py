import io
import os
import tempfile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph
from pdf2docx import Converter

_TEMPLATE_PDF = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Or_CV.pdf")


def _walk_paragraphs(elem, doc):
    """Yield all paragraphs in XML tree order, descending into tables."""
    tag = elem.tag
    if tag == qn("w:p"):
        yield DocxParagraph(elem, doc)
    elif tag in (qn("w:tbl"), qn("w:tr"), qn("w:tc")):
        for child in elem:
            yield from _walk_paragraphs(child, doc)


def _iter_all_paragraphs(doc: Document):
    for child in doc.element.body:
        yield from _walk_paragraphs(child, doc)


def _set_paragraph_text(para: DocxParagraph, text: str):
    """Replace paragraph text, preserving the first run's character formatting."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def build_docx_from_pdf_template(tailored_cv_text: str, cover_letter_text: str) -> bytes:
    if not os.path.exists(_TEMPLATE_PDF):
        raise FileNotFoundError(
            f"Template PDF not found at {_TEMPLATE_PDF!r}. "
            "Place Or_CV.pdf in the project root."
        )

    # Step 1: Convert PDF → docx, preserving fonts, colours, and layout
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        converter = Converter(_TEMPLATE_PDF)
        converter.convert(tmp_path)
        converter.close()
        doc = Document(tmp_path)
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # Step 2: Walk every paragraph in the converted document (incl. inside tables)
    all_paras = list(_iter_all_paragraphs(doc))
    new_lines = tailored_cv_text.splitlines()

    # Step 3: Overwrite text line-by-line, keeping each paragraph's original formatting
    for i, para in enumerate(all_paras):
        _set_paragraph_text(para, new_lines[i] if i < len(new_lines) else "")

    # Step 4: If tailored CV is longer than the original, append the extra lines
    for extra_line in new_lines[len(all_paras):]:
        doc.add_paragraph(extra_line)

    # Step 5: Append cover letter on a fresh page
    doc.add_page_break()
    doc.add_heading("Cover Letter", level=1)
    for line in cover_letter_text.splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
